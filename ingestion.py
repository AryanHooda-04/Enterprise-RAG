from __future__ import annotations

import base64
import hashlib
import io
import logging
import mimetypes
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import tiktoken
from openai import BadRequestError

from config import Settings, settings
from errors import DocumentParsingError, EmbeddingError, clean_external_error
from openai_client import get_openai_client
from usage_store import record_openai_usage
from vector_store import VectorStore


logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class TextSegment:
    text: str
    file_name: str
    page_number: int | None = None
    source_type: str = "text"
    image_index: int | None = None


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}


def suppress_mupdf_diagnostics(fitz_module) -> None:
    tools = getattr(fitz_module, "TOOLS", None)
    if tools is None:
        return
    for function_name in ("mupdf_display_errors", "mupdf_display_warnings"):
        function = getattr(tools, function_name, None)
        if callable(function):
            try:
                function(False)
            except Exception:
                pass


def safe_filename(filename: str | None) -> str:
    base = Path(filename or "document").name
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", base).strip("._")
    return cleaned or "document"


def compute_sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def compute_file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def load_document(file_path: Path, active_settings: Settings = settings) -> list[TextSegment]:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path, active_settings)
    if suffix == ".txt":
        return _load_txt(file_path)
    if suffix == ".docx":
        return _load_docx(file_path, active_settings)
    if suffix in IMAGE_EXTENSIONS:
        return _load_image(file_path, active_settings)

    raise DocumentParsingError(f"Unsupported file type: {suffix}")


def chunk_document(
    segments: list[TextSegment],
    *,
    file_hash: str,
    chunk_size_tokens: int | None = None,
    chunk_overlap_tokens: int | None = None,
    active_settings: Settings = settings,
) -> list[dict[str, Any]]:
    chunk_size = chunk_size_tokens or active_settings.chunk_size_tokens
    overlap = chunk_overlap_tokens if chunk_overlap_tokens is not None else active_settings.chunk_overlap_tokens

    if chunk_size <= 0:
        raise DocumentParsingError("chunk_size_tokens must be greater than zero.")
    if overlap < 0:
        raise DocumentParsingError("chunk_overlap_tokens cannot be negative.")
    if overlap >= chunk_size:
        raise DocumentParsingError("chunk_overlap_tokens must be smaller than chunk_size_tokens.")

    tokenizer = _get_tokenizer(active_settings.openai_embedding_model)
    chunks: list[dict[str, Any]] = []
    chunk_index = 0

    for segment in segments:
        text = segment.text.strip()
        if not text:
            continue

        tokens = tokenizer.encode(text)
        start = 0

        while start < len(tokens):
            end = min(start + chunk_size, len(tokens))
            chunk_text = tokenizer.decode(tokens[start:end]).strip()

            if chunk_text:
                metadata = {
                    "chunk_id": f"{file_hash}:{chunk_index}",
                    "file_hash": file_hash,
                    "file_name": segment.file_name,
                    "page_number": segment.page_number,
                    "source_type": segment.source_type,
                    "image_index": segment.image_index,
                    "chunk_index": chunk_index,
                    "token_start": start,
                    "token_count": end - start,
                }
                chunks.append({"text": chunk_text, "metadata": metadata})
                chunk_index += 1

            if end >= len(tokens):
                break
            start = end - overlap

    if not chunks:
        raise DocumentParsingError("No extractable text was found in the document.")

    return chunks


def generate_embeddings(
    texts: list[str],
    *,
    active_settings: Settings = settings,
) -> list[list[float]]:
    cleaned_texts = [text.strip() for text in texts if text and text.strip()]
    if not cleaned_texts:
        raise EmbeddingError("No text was provided for embedding generation.")

    client = get_openai_client()
    embeddings: list[list[float]] = []

    try:
        for start in range(0, len(cleaned_texts), active_settings.embedding_batch_size):
            batch = cleaned_texts[start : start + active_settings.embedding_batch_size]
            logger.info("Generating embeddings for batch %s-%s", start + 1, start + len(batch))
            response = client.embeddings.create(
                model=active_settings.openai_embedding_model,
                input=batch,
            )
            record_openai_usage(
                "embedding",
                active_settings.openai_embedding_model,
                getattr(response, "usage", None),
                active_settings=active_settings,
                input_count=len(batch),
            )
            batch_embeddings = [
                item.embedding for item in sorted(response.data, key=lambda item: item.index)
            ]
            embeddings.extend(batch_embeddings)
    except Exception as exc:
        raise EmbeddingError(clean_external_error(exc, "Embedding generation")) from exc

    logger.info("Generated %s embeddings", len(embeddings))
    return embeddings


def ingest_file(
    file_path: Path,
    vector_store: VectorStore,
    *,
    file_hash: str | None = None,
    display_name: str | None = None,
    chunk_size_tokens: int | None = None,
    chunk_overlap_tokens: int | None = None,
    active_settings: Settings = settings,
) -> dict[str, Any]:
    resolved_hash = file_hash or compute_file_hash(file_path)
    file_name = display_name or file_path.name

    existing_document = vector_store.get_document(resolved_hash)
    if existing_document:
        logger.info("Skipping duplicate upload for %s", file_name)
        return {
            "file_name": existing_document["file_name"],
            "file_hash": resolved_hash,
            "chunks_added": 0,
            "total_chunks": existing_document["chunk_count"],
            "skipped": True,
        }

    segments = [
        TextSegment(
            text=segment.text,
            file_name=file_name,
            page_number=segment.page_number,
            source_type=segment.source_type,
            image_index=segment.image_index,
        )
        for segment in load_document(file_path, active_settings)
    ]
    chunks = chunk_document(
        segments,
        file_hash=resolved_hash,
        chunk_size_tokens=chunk_size_tokens,
        chunk_overlap_tokens=chunk_overlap_tokens,
        active_settings=active_settings,
    )
    embeddings = generate_embeddings([chunk["text"] for chunk in chunks], active_settings=active_settings)
    added = vector_store.add_chunks(
        chunks,
        embeddings,
        file_hash=resolved_hash,
        file_name=file_name,
        source_path=str(file_path.resolve()),
    )
    vector_store.save()

    logger.info("Upload success for %s with %s chunks", file_name, added)
    return {
        "file_name": file_name,
        "file_hash": resolved_hash,
        "chunks_added": added,
        "total_chunks": added,
        "skipped": False,
    }


def _load_pdf(file_path: Path, active_settings: Settings = settings) -> list[TextSegment]:
    errors: list[str] = []
    segments: list[TextSegment] = []

    try:
        import pdfplumber

        with pdfplumber.open(str(file_path)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    segments.append(
                        TextSegment(text=text, file_name=file_path.name, page_number=page_number)
                    )
    except Exception as exc:
        errors.append(f"pdfplumber: {exc}")

    if not segments:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            for page_number, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    segments.append(TextSegment(text=text, file_name=file_path.name, page_number=page_number))
        except Exception as exc:
            errors.append(f"pypdf: {exc}")

    if active_settings.vision_ingestion_enabled:
        segments.extend(_load_pdf_visual_segments(file_path, active_settings))

    if segments:
        return segments

    detail = "; ".join(errors) if errors else "No parser returned text."
    raise DocumentParsingError(f"Unable to extract text from PDF. {detail}")


def _load_txt(file_path: Path) -> list[TextSegment]:
    last_error: Exception | None = None
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = file_path.read_text(encoding=encoding)
            if not text.strip():
                raise DocumentParsingError("TXT file is empty.")
            return [TextSegment(text=text, file_name=file_path.name)]
        except UnicodeDecodeError as exc:
            last_error = exc

    raise DocumentParsingError(f"Unable to decode TXT file. {last_error}")


def _load_image(file_path: Path, active_settings: Settings) -> list[TextSegment]:
    if not active_settings.vision_ingestion_enabled:
        raise DocumentParsingError("Image understanding is disabled for ingestion.")

    image_bytes = file_path.read_bytes()
    description = _describe_image_bytes(
        image_bytes,
        _mime_type_for_path(file_path),
        active_settings,
        context_label=file_path.name,
    )
    if not description:
        raise DocumentParsingError("No meaningful visual content was found in the image.")

    return [
        TextSegment(
            text=f"Visual content from {file_path.name}:\n{description}",
            file_name=file_path.name,
            source_type="image",
            image_index=1,
        )
    ]


def _load_pdf_visual_segments(file_path: Path, active_settings: Settings) -> list[TextSegment]:
    try:
        import fitz
    except Exception as exc:
        logger.info("PyMuPDF is unavailable for PDF page rendering; using embedded image fallback: %s", exc)
        return _load_pdf_embedded_image_segments(file_path, active_settings)

    suppress_mupdf_diagnostics(fitz)
    segments: list[TextSegment] = []
    visual_count = 0

    try:
        with fitz.open(str(file_path)) as document:
            for page_index in range(len(document)):
                if visual_count >= active_settings.max_visual_pages:
                    break

                page = document[page_index]
                page_number = page_index + 1
                page_text = (page.get_text("text") or "").strip()
                has_images = bool(page.get_images(full=True))

                if page_text and not has_images:
                    continue

                try:
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    image_bytes = pixmap.tobytes("png")
                except Exception as exc:
                    logger.warning("Unable to render PDF page %s from %s: %s", page_number, file_path.name, exc)
                    continue

                visual_count += 1
                description = _describe_image_bytes(
                    image_bytes,
                    "image/png",
                    active_settings,
                    context_label=f"{file_path.name}, page {page_number}",
                )
                if description:
                    segments.append(
                        TextSegment(
                            text=f"Visual content from {file_path.name}, page {page_number}:\n{description}",
                            file_name=file_path.name,
                            page_number=page_number,
                            source_type="image",
                            image_index=visual_count,
                        )
                    )
    except DocumentParsingError:
        raise
    except Exception as exc:
        logger.warning("Unable to render PDF visual pages from %s: %s", file_path.name, exc)
        return _load_pdf_embedded_image_segments(file_path, active_settings)

    if segments:
        logger.info("Extracted %s visual segment(s) from %s", len(segments), file_path.name)
    return segments


def _load_pdf_embedded_image_segments(file_path: Path, active_settings: Settings) -> list[TextSegment]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        segments: list[TextSegment] = []
        image_count = 0

        for page_number, page in enumerate(reader.pages, start=1):
            for image in getattr(page, "images", []):
                if image_count >= active_settings.max_visual_pages:
                    break

                image_count += 1
                image_name = getattr(image, "name", f"image_{image_count}.jpg")
                description = _describe_image_bytes(
                    image.data,
                    _mime_type_for_path(Path(image_name)),
                    active_settings,
                    context_label=f"{file_path.name}, page {page_number}, embedded image {image_count}",
                )
                if description:
                    segments.append(
                        TextSegment(
                            text=(
                                f"Visual content from {file_path.name}, page {page_number}, "
                                f"embedded image {image_count}:\n{description}"
                            ),
                            file_name=file_path.name,
                            page_number=page_number,
                            source_type="image",
                            image_index=image_count,
                        )
                    )

            if image_count >= active_settings.max_visual_pages:
                break

        if segments:
            logger.info("Extracted %s embedded PDF image segment(s) from %s", len(segments), file_path.name)
        return segments
    except DocumentParsingError:
        raise
    except Exception as exc:
        raise DocumentParsingError(f"Unable to extract embedded images from PDF: {exc}") from exc


def _load_docx_visual_segments(file_path: Path, active_settings: Settings) -> list[TextSegment]:
    segments: list[TextSegment] = []
    image_count = 0

    try:
        with zipfile.ZipFile(file_path) as archive:
            media_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("word/media/")
                and Path(name).suffix.lower() in IMAGE_EXTENSIONS
            )

            for media_name in media_names[: active_settings.max_docx_images]:
                image_count += 1
                description = _describe_image_bytes(
                    archive.read(media_name),
                    _mime_type_for_path(Path(media_name)),
                    active_settings,
                    context_label=f"{file_path.name}, embedded image {image_count}",
                )
                if description:
                    segments.append(
                        TextSegment(
                            text=(
                                f"Visual content from {file_path.name}, embedded image "
                                f"{image_count}:\n{description}"
                            ),
                            file_name=file_path.name,
                            source_type="image",
                            image_index=image_count,
                        )
                    )
    except DocumentParsingError:
        raise
    except Exception as exc:
        raise DocumentParsingError(f"Unable to extract visual content from DOCX: {exc}") from exc

    if segments:
        logger.info("Extracted %s embedded image segment(s) from %s", len(segments), file_path.name)
    return segments


def _load_docx(file_path: Path, active_settings: Settings = settings) -> list[TextSegment]:
    try:
        from docx import Document

        document = Document(str(file_path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        segments: list[TextSegment] = []
        text = "\n".join(parts)
        if text.strip():
            segments.append(TextSegment(text=text, file_name=file_path.name))

        if active_settings.vision_ingestion_enabled:
            segments.extend(_load_docx_visual_segments(file_path, active_settings))

        if not segments:
            raise DocumentParsingError("DOCX file contains no extractable text or images.")

        return segments
    except DocumentParsingError:
        raise
    except Exception as exc:
        raise DocumentParsingError(f"Unable to extract text from DOCX: {exc}") from exc


def _describe_image_bytes(
    image_bytes: bytes,
    mime_type: str,
    active_settings: Settings,
    *,
    context_label: str,
) -> str:
    if not image_bytes:
        return ""

    data_url = _image_bytes_to_data_url(image_bytes, mime_type, active_settings.max_image_dimension_px)
    detail = _vision_detail(active_settings.vision_detail)
    prompt = (
        "Prepare this image for enterprise document retrieval. Extract all readable text, "
        "headings, labels, table values, chart trends, UI fields, numbers, and named entities. "
        "Describe diagrams, screenshots, and relationships that would help semantic search. "
        "Only include visible information. If the image has no useful document content, say "
        "'No meaningful visual content'. "
        f"Source: {context_label}."
    )

    client = get_openai_client()
    params: dict[str, Any] = {
        "model": active_settings.openai_vision_model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": data_url,
                            "detail": detail,
                        },
                    },
                ],
            }
        ],
        "max_completion_tokens": 700,
    }
    if active_settings.openai_temperature is not None:
        params["temperature"] = active_settings.openai_temperature

    try:
        response = _create_chat_completion_with_fallbacks(client, params)
        record_openai_usage(
            "vision",
            active_settings.openai_vision_model,
            getattr(response, "usage", None),
            active_settings=active_settings,
            document_name=context_label,
            input_count=1,
        )
        description = (response.choices[0].message.content or "").strip()
    except Exception as exc:
        raise DocumentParsingError(clean_external_error(exc, "Image understanding")) from exc

    if description.lower().startswith("no meaningful visual content"):
        return ""
    logger.info("Generated visual description for %s", context_label)
    return description


def _image_bytes_to_data_url(image_bytes: bytes, mime_type: str, max_dimension_px: int) -> str:
    try:
        from PIL import Image, ImageOps

        with Image.open(io.BytesIO(image_bytes)) as image:
            image = ImageOps.exif_transpose(image)
            image.seek(0)
            if image.mode in {"RGBA", "LA", "P"}:
                image = image.convert("RGBA")
                background = Image.new("RGB", image.size, "white")
                background.paste(image, mask=image.getchannel("A"))
                image = background
            else:
                image = image.convert("RGB")

            max_dimension = max(256, int(max_dimension_px or 1600))
            image.thumbnail((max_dimension, max_dimension), Image.Resampling.LANCZOS)
            output = io.BytesIO()
            image.save(output, format="JPEG", quality=88, optimize=True)
            encoded = base64.b64encode(output.getvalue()).decode("ascii")
            return f"data:image/jpeg;base64,{encoded}"
    except Exception as exc:
        logger.warning("Unable to normalize image before vision request: %s", exc)

    encoded = base64.b64encode(image_bytes).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _create_chat_completion_with_fallbacks(client, params: dict[str, Any]):
    last_error: Exception | None = None

    for _ in range(3):
        try:
            return client.chat.completions.create(**params)
        except BadRequestError as exc:
            last_error = exc
            message = str(exc).lower()
            changed = False

            if "max_completion_tokens" in message and "max_tokens" not in params:
                params["max_tokens"] = params.pop("max_completion_tokens")
                changed = True

            if "temperature" in message and "temperature" in params:
                params.pop("temperature")
                changed = True

            if not changed:
                raise

    if last_error:
        raise last_error
    raise DocumentParsingError("Image understanding failed.")


def _mime_type_for_path(file_path: Path) -> str:
    mime_type, _ = mimetypes.guess_type(file_path.name)
    return mime_type or "image/jpeg"


def _vision_detail(value: str) -> str:
    detail = (value or "auto").strip().lower()
    if detail in {"low", "high", "auto"}:
        return detail
    return "auto"


def _get_tokenizer(model_name: str):
    try:
        return tiktoken.encoding_for_model(model_name)
    except KeyError:
        return tiktoken.get_encoding("cl100k_base")
